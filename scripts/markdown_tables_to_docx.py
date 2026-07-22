from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_east_asia_font(run, font_name="Microsoft JhengHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=8):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text.strip())
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.font.bold = bold


def split_table_row(line):
    row = line.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [cell.strip().replace("<br>", "\n") for cell in row.split("|")]


def is_table_separator(line):
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in cells)


def is_table_start(lines, index):
    return (
        index + 1 < len(lines)
        and lines[index].strip().startswith("|")
        and lines[index + 1].strip().startswith("|")
        and is_table_separator(lines[index + 1])
    )


def add_markdown_table(doc, table_lines):
    header = split_table_row(table_lines[0])
    body = [split_table_row(line) for line in table_lines[2:]]
    col_count = len(header)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, header[idx] if idx < len(header) else "", bold=True, size=8)

    for row_values in body:
        row = table.add_row()
        for idx, cell in enumerate(row.cells):
            text = row_values[idx] if idx < len(row_values) else ""
            set_cell_text(cell, text, size=8)

    doc.add_paragraph()


def add_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    return paragraph


def convert(md_path, docx_path):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if is_table_start(lines, i):
            table_lines = [lines[i], lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped[2:].strip())
            set_east_asia_font(run)
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue

        if stripped.startswith("## "):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(stripped[3:].strip())
            set_east_asia_font(run)
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(stripped[4:].strip())
            set_east_asia_font(run)
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x5B, 0x2C, 0x06)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("─" * 70)
            i += 1
            continue

        if stripped.startswith("> "):
            paragraph = add_paragraph(doc, stripped[2:].strip())
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(stripped[2:].strip())
            set_east_asia_font(run)
            run.font.size = Pt(10.5)
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    doc.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: markdown_tables_to_docx.py input.md output.docx")
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
